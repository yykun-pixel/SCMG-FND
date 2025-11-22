import sys
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    print("[Error] python-docx is not installed. Please install with: pip install python-docx")
    sys.exit(1)

def docx_to_markdown(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    lines = []

    # 处理段落（粗略识别标题）
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            lines.append("")
            continue
        style_name = getattr(p.style, 'name', '').lower()
        if 'heading' in style_name:
            # 将 Heading 级别映射为 Markdown #
            level = 1
            for k in ['heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5', 'heading 6']:
                if k in style_name:
                    level = int(k.split()[-1])
                    break
            level = max(1, min(level, 6))
            lines.append('#' * level + ' ' + text)
        else:
            lines.append(text)

    # 处理表格
    for tbl in doc.tables:
        max_cols = max(len(row.cells) for row in tbl.rows) if tbl.rows else 0
        if max_cols == 0:
            continue
        # 表头（第一行）
        header = [tbl.rows[0].cells[i].text.strip() if i < len(tbl.rows[0].cells) else '' for i in range(max_cols)]
        lines.append('')
        lines.append('| ' + ' | '.join(h or ' ' for h in header) + ' |')
        lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
        # 其余行
        for r in tbl.rows[1:]:
            row_vals = [r.cells[i].text.strip() if i < len(r.cells) else '' for i in range(max_cols)]
            lines.append('| ' + ' | '.join(row_vals) + ' |')
        lines.append('')

    return '\n'.join(lines)


def main():
    if len(sys.argv) < 3:
        print("Usage: python scripts/convert_docx_to_text.py <input.docx> <output.md>")
        sys.exit(1)

    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])

    if not in_path.exists():
        print(f"[Error] Input file not found: {in_path}")
        sys.exit(1)

    md_text = docx_to_markdown(in_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(md_text, encoding='utf-8')
    print(f"[OK] Converted to: {out_path}")

if __name__ == "__main__":
    main() 